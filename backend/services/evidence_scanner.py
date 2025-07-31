"""
Evidence Scanner - Professional evidence analysis and classification
"""
import io
import re
from datetime import UTC, datetime
from email.parser import Parser
from pathlib import Path
from typing import Any

import docx  # type: ignore[import-untyped]
import fitz  # type: ignore[import-untyped]  # PyMuPDF for PDF processing
import pytesseract  # type: ignore[import-untyped]
from PIL import Image

from backend.models.document import Document, DocumentType


class EvidenceType:
    """Evidence classification types"""

    CONTRACT = "contract"
    CORRESPONDENCE = "correspondence"
    FINANCIAL = "financial_record"
    WITNESS_STATEMENT = "witness_statement"
    EXPERT_REPORT = "expert_report"
    COURT_DOCUMENT = "court_document"
    PHOTO_VIDEO = "photo_video_evidence"
    MEDICAL_RECORD = "medical_record"
    POLICE_REPORT = "police_report"
    BUSINESS_RECORD = "business_record"


class EvidenceScanner:
    """Professional evidence scanner for UK legal practice"""

    # Key patterns for evidence classification
    EVIDENCE_PATTERNS = {
        EvidenceType.CONTRACT: [
            r"agreement",
            r"contract",
            r"terms and conditions",
            r"deed",
            r"whereas",
            r"party of the first part",
            r"consideration",
            r"covenant",
            r"warranty",
            r"indemnity",
        ],
        EvidenceType.CORRESPONDENCE: [
            r"dear",
            r"yours sincerely",
            r"yours faithfully",
            r"kind regards",
            r"please find",
            r"further to",
            r"we write",
            r"letter before action",
        ],
        EvidenceType.FINANCIAL: [
            r"invoice",
            r"receipt",
            r"payment",
            r"balance sheet",
            r"profit and loss",
            r"bank statement",
            r"transaction",
            r"£\d+",
            r"total due",
        ],
        EvidenceType.WITNESS_STATEMENT: [
            r"i [name] make this statement",
            r"statement of truth",
            r"i believe.*facts.*true",
            r"witness statement",
            r"i saw",
            r"i heard",
        ],
        EvidenceType.COURT_DOCUMENT: [
            r"claim form",
            r"particulars of claim",
            r"defence",
            r"court order",
            r"judgment",
            r"application notice",
            r"witness summons",
            r"bundle",
        ],
        EvidenceType.MEDICAL_RECORD: [
            r"diagnosis",
            r"prognosis",
            r"medical report",
            r"patient",
            r"treatment",
            r"injury",
            r"examination",
            r"dr\.",
            r"consultant",
        ],
    }

    # Legal significance indicators
    SIGNIFICANCE_KEYWORDS = {
        "high": [
            "breach",
            "default",
            "negligence",
            "liability",
            "damages",
            "termination",
            "injury",
            "loss",
            "fraud",
            "misrepresentation",
        ],
        "medium": [
            "dispute",
            "complaint",
            "concern",
            "issue",
            "problem",
            "disagreement",
            "claim",
            "allegation",
        ],
        "low": ["information", "update", "notice", "acknowledgment"],
    }

    def __init__(self):
        self.ocr_available = self._check_ocr_availability()

    def _check_ocr_availability(self) -> bool:
        """Check if OCR tools are available"""
        try:
            pytesseract.get_tesseract_version()
            return True
        except Exception:
            return False

    async def scan_document(self, document_path: Path, document: Document) -> dict[str, Any]:
        """Comprehensive document scanning and analysis"""
        scan_result: dict[str, Any] = {
            "document_id": str(document.id),
            "scan_date": datetime.now(UTC).isoformat(),
            "file_type": document_path.suffix.lower(),
            "evidence_type": None,
            "significance": "low",
            "key_information": {},
            "extracted_data": {},
            "legal_relevance": {},
            "authenticity_check": {},
            "chain_of_custody": [],
            "admissibility_issues": [],
        }

        # Extract text based on file type
        file_type = scan_result["file_type"]
        extracted_text = await self._extract_text(document_path, file_type)
        scan_result["extracted_text_preview"] = extracted_text[:500] if extracted_text else ""

        # Classify evidence type
        scan_result["evidence_type"] = self._classify_evidence(extracted_text, document)

        # Extract key information
        scan_result["key_information"] = self._extract_key_info(extracted_text)

        # Assess legal significance
        scan_result["significance"] = self._assess_significance(extracted_text)

        # Check authenticity markers
        scan_result["authenticity_check"] = self._check_authenticity(document_path, extracted_text)

        # Identify admissibility issues
        scan_result["admissibility_issues"] = self._check_admissibility(scan_result)

        # Extract structured data
        evidence_type = scan_result["evidence_type"]
        scan_result["extracted_data"] = self._extract_structured_data(extracted_text, evidence_type)

        # Legal relevance assessment
        scan_result["legal_relevance"] = self._assess_legal_relevance(extracted_text, evidence_type)

        return scan_result

    async def _extract_text(self, file_path: Path, file_type: str) -> str:
        """Extract text from various file formats"""
        try:
            if file_type == ".pdf":
                return self._extract_pdf_text(file_path)
            if file_type in [".doc", ".docx"]:
                return self._extract_word_text(file_path)
            if file_type in [".jpg", ".jpeg", ".png", ".tiff"]:
                return self._extract_image_text(file_path)
            if file_type == ".eml":
                return self._extract_email_text(file_path)
            if file_type in [".txt", ".rtf"]:
                return file_path.read_text(encoding="utf-8", errors="ignore")
            return ""
        except Exception:
            return ""

    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF with metadata"""
        text: list[str] = []
        try:
            with fitz.open(str(file_path)) as doc:  # type: ignore
                # Extract metadata
                metadata = doc.metadata  # type: ignore[attr-defined]
                if metadata:
                    text.append(f"PDF Metadata: Created {metadata.get('creationDate', 'Unknown')}")  # type: ignore[union-attr]
                    text.append(f"Author: {metadata.get('author', 'Unknown')}")  # type: ignore[union-attr]

                # Extract text from each page
                for page_num, page in enumerate(doc, 1):  # type: ignore[arg-type]
                    page_text = page.get_text()  # type: ignore[attr-defined]
                    if page_text.strip():  # type: ignore[union-attr]
                        text.append(f"\n--- Page {page_num} ---\n{page_text}")

                    # If no text, try OCR on the page
                    if not page_text.strip() and self.ocr_available:  # type: ignore[union-attr]
                        pix = page.get_pixmap()  # type: ignore[attr-defined]
                        img_data = pix.pil_tobytes(format="PNG")  # type: ignore[attr-defined]
                        ocr_text = self._ocr_image_bytes(img_data)  # type: ignore[arg-type]
                        if ocr_text:
                            text.append(f"\n--- Page {page_num} (OCR) ---\n{ocr_text}")

        except Exception:
            pass

        return "\n".join(text)

    def _extract_word_text(self, file_path: Path) -> str:
        """Extract text from Word documents"""
        try:
            doc = docx.Document(str(file_path))
            text: list[str] = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text: list[str] = [cell.text.strip() for cell in row.cells]
                    if any(row_text):
                        text.append(" | ".join(row_text))

            return "\n".join(text)
        except Exception:
            return ""

    def _extract_image_text(self, file_path: Path) -> str:
        """Extract text from images using OCR"""
        if not self.ocr_available:
            return ""

        try:
            image = Image.open(file_path)  # type: ignore[no-untyped-call]
            text_result = pytesseract.image_to_string(image)  # type: ignore[no-any-return]
            text = str(text_result) if isinstance(text_result, str) else ""

            # Also extract image metadata
            exif_data = image.getexif()  # type: ignore[no-untyped-call]
            if exif_data:
                text = f"Image metadata: {exif_data}\n\n{text}"

            return text
        except Exception:
            return ""

    def _extract_email_text(self, file_path: Path) -> str:
        """Extract text from email files"""
        try:
            with open(file_path, encoding="utf-8", errors="ignore") as f:
                msg = Parser().parse(f)

            text: list[str] = []
            text.append(f"From: {msg.get('From', 'Unknown')}")
            text.append(f"To: {msg.get('To', 'Unknown')}")
            text.append(f"Date: {msg.get('Date', 'Unknown')}")
            text.append(f"Subject: {msg.get('Subject', 'Unknown')}")
            text.append("\nBody:")

            # Extract body
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        payload = part.get_payload(decode=True)
                        if payload and isinstance(payload, bytes):
                            text.append(payload.decode("utf-8", errors="ignore"))
                        elif payload and isinstance(payload, str):
                            text.append(payload)
            else:
                payload = msg.get_payload(decode=True)
                if payload and isinstance(payload, bytes):
                    text.append(payload.decode("utf-8", errors="ignore"))
                elif payload and isinstance(payload, str):
                    text.append(payload)

            return "\n".join(text)
        except Exception:
            return ""

    def _ocr_image_bytes(self, image_bytes: bytes) -> str:
        """OCR from image bytes"""
        try:
            image = Image.open(io.BytesIO(image_bytes))  # type: ignore[no-untyped-call]
            result = pytesseract.image_to_string(image)  # type: ignore[no-any-return]
            return str(result) if isinstance(result, str) else ""
        except Exception:
            return ""

    def _classify_evidence(self, text: str, document: Document) -> str:
        """Classify evidence type using patterns and context"""
        if not text:
            return EvidenceType.BUSINESS_RECORD

        text_lower = text.lower()
        scores = {}

        # Score each evidence type based on pattern matches
        for evidence_type, patterns in self.EVIDENCE_PATTERNS.items():
            score = 0
            for pattern in patterns:
                if re.search(pattern, text_lower):
                    score += 1
            scores[evidence_type] = score

        # Check document metadata
        if document.document_type.value == DocumentType.CONTRACT.value:
            scores[EvidenceType.CONTRACT] += 5
        elif document.document_type.value == DocumentType.COURT_FILING.value:
            scores[EvidenceType.COURT_DOCUMENT] += 5
        elif document.document_type.value == DocumentType.CORRESPONDENCE.value:
            scores[EvidenceType.CORRESPONDENCE] += 5

        # Return highest scoring type
        if scores:
            return max(scores, key=scores.get)  # type: ignore[arg-type]
        return EvidenceType.BUSINESS_RECORD

    def _extract_key_info(self, text: str) -> dict[str, list[str]]:
        """Extract key information from text"""
        if not text:
            return {}

        key_info: dict[str, list[str]] = {
            "dates": [],
            "amounts": [],
            "names": [],
            "addresses": [],
            "reference_numbers": [],
            "email_addresses": [],
            "phone_numbers": [],
        }

        # Date patterns (UK format)
        date_patterns = [
            r"\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b",
            r"\b(\d{1,2}\s+(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{2,4})\b",
            r"\b((?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{2,4})\b",
        ]

        for pattern in date_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            key_info["dates"].extend(matches)

        # Amounts (UK currency)
        amount_pattern = r"£[\d,]+(?:\.\d{2})?"
        key_info["amounts"] = re.findall(amount_pattern, text)

        # Email addresses
        email_pattern = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b"
        key_info["email_addresses"] = re.findall(email_pattern, text)

        # Phone numbers (UK format)
        phone_patterns = [
            r"\b(?:0\d{4}\s?\d{6})\b",  # UK landline
            r"\b(?:07\d{3}\s?\d{6})\b",  # UK mobile
            r"\b(?:\+44\s?7\d{3}\s?\d{6})\b",  # International UK mobile
        ]

        for pattern in phone_patterns:
            matches = re.findall(pattern, text)
            key_info["phone_numbers"].extend(matches)

        # Reference numbers (case numbers, invoice numbers, etc.)
        ref_patterns = [
            r"\b(?:Case|Ref|Reference|Invoice|Order|Claim)\s*(?:No|Number|#)?:?\s*([A-Z0-9-/]+)\b",
            r"\b([A-Z]{2,}-\d{4}-\d+)\b",  # Pattern like CASE-2024-001
        ]

        for pattern in ref_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            key_info["reference_numbers"].extend(matches)

        # Names (basic extraction - would use NER in production)
        # Look for patterns like "Mr/Mrs/Ms/Dr Name"
        name_pattern = r"\b(?:Mr|Mrs|Ms|Dr|Prof)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\b"
        key_info["names"] = re.findall(name_pattern, text)

        # Remove duplicates
        for key in key_info:
            key_info[key] = list(set(key_info[key]))

        return key_info

    def _assess_significance(self, text: str) -> str:
        """Assess legal significance of the evidence"""
        if not text:
            return "low"

        text_lower = text.lower()

        # Count significance indicators
        high_count = sum(1 for keyword in self.SIGNIFICANCE_KEYWORDS["high"] if keyword in text_lower)
        medium_count = sum(1 for keyword in self.SIGNIFICANCE_KEYWORDS["medium"] if keyword in text_lower)

        if high_count >= 3:
            return "high"
        if high_count >= 1 or medium_count >= 3:
            return "medium"
        return "low"

    def _check_authenticity(self, file_path: Path, text: str) -> dict[str, Any]:
        """Check document authenticity markers"""
        authenticity: dict[str, Any] = {
            "digital_signature": False,
            "metadata_intact": True,
            "creation_date": None,
            "modification_date": None,
            "author": None,
            "suspicious_markers": [],
            "authenticity_score": 100,
        }

        # Check file metadata
        try:
            stat = file_path.stat()
            authenticity["creation_date"] = datetime.fromtimestamp(stat.st_ctime).isoformat()
            authenticity["modification_date"] = datetime.fromtimestamp(stat.st_mtime).isoformat()

            # Check if modified after creation (suspicious for evidence)
            if stat.st_mtime > stat.st_ctime + 86400:  # Modified more than 1 day after creation
                authenticity["suspicious_markers"].append("File modified significantly after creation")
                authenticity["authenticity_score"] -= 20

        except Exception:
            authenticity["metadata_intact"] = False
            authenticity["authenticity_score"] -= 30

        # Check for PDF signatures
        if file_path.suffix.lower() == ".pdf":
            try:
                with fitz.open(str(file_path)) as doc:  # type: ignore
                    for page in doc:  # type: ignore
                        if page.get_text().find("/Sig") != -1:  # type: ignore[attr-defined]
                            authenticity["digital_signature"] = True
                            authenticity["authenticity_score"] += 10
                            break
            except Exception:
                pass

        # Check for suspicious content patterns
        suspicious_patterns = [
            (r"edited\s+with\s+\w+", "Editing software marker found"),
            (r"screenshot", "Screenshot indicator found"),
            (r"copy\s+of\s+copy", "Multiple copy indicator"),
            (r"\[modified\]", "Modification marker found"),
        ]

        for pattern, message in suspicious_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                authenticity["suspicious_markers"].append(message)
                authenticity["authenticity_score"] -= 10

        # Ensure score is between 0 and 100
        authenticity["authenticity_score"] = max(0, min(100, authenticity["authenticity_score"]))

        return authenticity

    def _check_admissibility(self, scan_result: dict[str, Any]) -> list[str]:
        """Check potential admissibility issues under UK law"""
        issues: list[str] = []

        # Check authenticity score
        if scan_result["authenticity_check"]["authenticity_score"] < 70:
            issues.append("Low authenticity score - may face challenge on authenticity grounds")

        # Check for hearsay
        if scan_result["evidence_type"] in [
            EvidenceType.CORRESPONDENCE,
            EvidenceType.BUSINESS_RECORD,
        ] and "i was told" in scan_result.get("extracted_text_preview", "").lower():
            issues.append("Potential hearsay - consider Civil Evidence Act 1995 notice requirements")

        # Check for privilege
        text_preview = scan_result.get("extracted_text_preview", "").lower()
        if "without prejudice" in text_preview:
            issues.append("Without prejudice communication - generally not admissible")
        elif "legally privileged" in text_preview or "legal advice" in text_preview:
            issues.append("May be subject to legal professional privilege")

        # Check for data protection issues
        if scan_result["key_information"].get("email_addresses") or scan_result["key_information"].get("phone_numbers"):
            issues.append("Contains personal data - ensure GDPR compliance")

        # Expert evidence requirements
        if scan_result["evidence_type"] == EvidenceType.EXPERT_REPORT:
            required_elements: list[str] = [
                "qualifications",
                "instructions",
                "statement of truth",
            ]
            text_lower = scan_result.get("extracted_text_preview", "").lower()
            missing: list[str] = [elem for elem in required_elements if elem not in text_lower]
            if missing:
                issues.append(f"Expert report may be missing: {', '.join(missing)} (CPR Part 35)")

        return issues

    def _extract_structured_data(self, text: str, evidence_type: str) -> dict[str, Any]:
        """Extract structured data based on evidence type"""
        structured_data = {}

        if evidence_type == EvidenceType.CONTRACT:
            structured_data = self._extract_contract_data(text)
        elif evidence_type == EvidenceType.FINANCIAL:
            structured_data = self._extract_financial_data(text)
        elif evidence_type == EvidenceType.WITNESS_STATEMENT:
            structured_data = self._extract_witness_data(text)
        elif evidence_type == EvidenceType.CORRESPONDENCE:
            structured_data = self._extract_correspondence_data(text)

        return structured_data

    def _extract_contract_data(self, text: str) -> dict[str, Any]:
        """Extract contract-specific data"""
        data: dict[str, Any] = {
            "parties": [],
            "effective_date": None,
            "termination_date": None,
            "consideration": None,
            "key_terms": [],
            "termination_clause": None,
            "dispute_resolution": None,
        }

        # Extract parties
        party_pattern = r"between\s+([^,]+?)\s+(?:and|AND)\s+([^,]+?)(?:\s+\(|,)"
        matches = re.findall(party_pattern, text, re.IGNORECASE)
        if matches:
            data["parties"] = [party.strip() for match in matches for party in match]

        # Extract dates
        date_patterns = {
            "effective_date": r"(?:effective|commencement)\s+date[:\s]+([^,\n]+)",
            "termination_date": r"(?:termination|expiry)\s+date[:\s]+([^,\n]+)",
        }

        for key, pattern in date_patterns.items():
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                data[key] = match.group(1).strip()

        # Extract consideration
        consideration_pattern = r"consideration\s+of\s+(£[\d,]+(?:\.\d{2})?)"
        match = re.search(consideration_pattern, text, re.IGNORECASE)
        if match:
            data["consideration"] = match.group(1)

        # Key terms (simplified)
        if "warranty" in text.lower():
            data["key_terms"].append("Contains warranties")
        if "indemnity" in text.lower():
            data["key_terms"].append("Contains indemnities")
        if "confidential" in text.lower():
            data["key_terms"].append("Contains confidentiality clause")

        return data

    def _extract_financial_data(self, text: str) -> dict[str, Any]:
        """Extract financial document data"""
        data: dict[str, Any] = {
            "total_amount": None,
            "line_items": [],
            "payment_terms": None,
            "due_date": None,
            "invoice_number": None,
            "vat_amount": None,
        }

        # Extract amounts
        amounts = re.findall(r"£([\d,]+(?:\.\d{2})?)", text)
        if amounts:
            # Assume largest amount is total
            amounts_numeric: list[float] = [float(a.replace(",", "")) for a in amounts]
            data["total_amount"] = f"£{max(amounts_numeric):,.2f}"

        # Extract invoice number
        invoice_pattern = r"invoice\s*(?:no|number|#)?[:\s]*([A-Z0-9-]+)"
        match = re.search(invoice_pattern, text, re.IGNORECASE)
        if match:
            data["invoice_number"] = match.group(1)

        # Extract VAT
        vat_pattern = r"VAT\s*(?:@\s*20%)?[:\s]*(£[\d,]+(?:\.\d{2})?)"
        match = re.search(vat_pattern, text, re.IGNORECASE)
        if match:
            data["vat_amount"] = match.group(1)

        return data

    def _extract_witness_data(self, text: str) -> dict[str, Any]:
        """Extract witness statement data"""
        data: dict[str, Any] = {
            "witness_name": None,
            "statement_date": None,
            "incidents_described": [],
            "statement_of_truth": False,
        }

        # Extract witness name
        name_pattern = r"I,?\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*),?\s+(?:make|state|declare)"
        match = re.search(name_pattern, text)
        if match:
            data["witness_name"] = match.group(1)

        # Check for statement of truth
        if "believe that the facts" in text.lower() and "true" in text.lower():
            data["statement_of_truth"] = True

        # Extract incident descriptions
        incident_patterns: list[str] = [
            r"I\s+saw\s+([^.]+)",
            r"I\s+witnessed\s+([^.]+)",
            r"I\s+observed\s+([^.]+)",
        ]

        for pattern in incident_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            data["incidents_described"].extend(matches[:3])  # Limit to first 3

        return data

    def _extract_correspondence_data(self, text: str) -> dict[str, Any]:
        """Extract correspondence data"""
        data: dict[str, Any] = {
            "sender": None,
            "recipient": None,
            "date": None,
            "subject": None,
            "is_without_prejudice": False,
            "is_letter_before_action": False,
        }

        # Check for without prejudice
        if "without prejudice" in text.lower():
            data["is_without_prejudice"] = True

        # Check for letter before action
        if "letter before action" in text.lower() or "letter of claim" in text.lower():
            data["is_letter_before_action"] = True

        # Extract sender/recipient
        from_pattern = r"From:\s*([^\n]+)"
        to_pattern = r"To:\s*([^\n]+)"

        match = re.search(from_pattern, text, re.IGNORECASE)
        if match:
            data["sender"] = match.group(1).strip()

        match = re.search(to_pattern, text, re.IGNORECASE)
        if match:
            data["recipient"] = match.group(1).strip()

        return data

    def _assess_legal_relevance(self, text: str, evidence_type: str) -> dict[str, Any]:
        """Assess legal relevance of the evidence"""
        relevance: dict[str, Any] = {
            "relevance_score": 0,
            "relevant_to_issues": [],
            "supports_claims": [],
            "contradicts_claims": [],
            "procedural_importance": None,
        }

        text_lower = text.lower()

        # Check relevance to common legal issues
        issue_keywords = {
            "breach_of_contract": [
                "breach",
                "failed to",
                "did not",
                "violation",
                "default",
            ],
            "negligence": [
                "negligent",
                "duty of care",
                "reasonable care",
                "foreseeable",
            ],
            "damages": ["loss", "damage", "compensation", "costs", "expenses"],
            "liability": ["liable", "responsible", "fault", "causation"],
        }

        for issue, keywords in issue_keywords.items():
            if any(keyword in text_lower for keyword in keywords):
                relevance["relevant_to_issues"].append(issue)
                relevance["relevance_score"] += 20

        # Check for admissions
        admission_patterns: list[str] = [
            "acknowledge",
            "admit",
            "accept",
            "agree that",
            "concede",
        ]

        if any(pattern in text_lower for pattern in admission_patterns):
            relevance["supports_claims"].append("Contains potential admissions")
            relevance["relevance_score"] += 30

        # Check for denials
        denial_patterns: list[str] = [
            "deny",
            "dispute",
            "disagree",
            "reject",
            "unfounded",
        ]

        if any(pattern in text_lower for pattern in denial_patterns):
            relevance["contradicts_claims"].append("Contains denials or disputes")
            relevance["relevance_score"] += 15

        # Procedural importance
        if evidence_type == EvidenceType.COURT_DOCUMENT:
            relevance["procedural_importance"] = "High - Court document"
            relevance["relevance_score"] += 25
        elif "letter before action" in text_lower:
            relevance["procedural_importance"] = "High - Pre-action protocol"
            relevance["relevance_score"] += 20

        # Cap relevance score at 100
        relevance["relevance_score"] = min(100, relevance["relevance_score"])

        return relevance

    async def generate_evidence_report(self, case_id: str, scan_results: list[dict[str, Any]]) -> dict[str, Any]:
        """Generate comprehensive evidence report for the case"""
        report: dict[str, Any] = {
            "case_id": case_id,
            "report_date": datetime.now(UTC).isoformat(),
            "total_documents_scanned": len(scan_results),
            "evidence_summary": {},
            "key_evidence": [],
            "missing_evidence": [],
            "admissibility_concerns": [],
            "chain_of_custody_status": "Maintained",
            "recommendations": [],
        }

        # Categorize evidence
        evidence_by_type: dict[str, int] = {}
        high_significance_docs: list[dict[str, Any]] = []
        admissibility_issues: list[dict[str, Any]] = []

        for result in scan_results:
            # Count by type
            ev_type = result.get("evidence_type", "unknown")
            evidence_by_type[ev_type] = evidence_by_type.get(ev_type, 0) + 1

            # Identify key evidence
            if result.get("significance") == "high" or result.get("legal_relevance", {}).get("relevance_score", 0) > 70:
                high_significance_docs.append(
                    {
                        "document_id": result["document_id"],
                        "type": ev_type,
                        "significance": result.get("significance"),
                        "key_points": result.get("key_information", {}),
                    }
                )

            # Collect admissibility issues
            if result.get("admissibility_issues"):
                admissibility_issues.extend(
                    [{"document_id": result["document_id"], "issue": issue} for issue in result["admissibility_issues"]]
                )

        report["evidence_summary"] = evidence_by_type
        report["key_evidence"] = high_significance_docs[:10]  # Top 10
        report["admissibility_concerns"] = admissibility_issues

        # Generate recommendations
        if not evidence_by_type.get(EvidenceType.CONTRACT) and "contract" in case_id.lower():
            report["recommendations"].append("Obtain original contract or agreement")

        if not evidence_by_type.get(EvidenceType.WITNESS_STATEMENT):
            report["recommendations"].append("Consider obtaining witness statements")

        if admissibility_issues:
            report["recommendations"].append("Address admissibility concerns before trial")

        return report
